from django.http import HttpResponse
from django.shortcuts import render
from django.contrib import messages
from django.core.files.storage import default_storage
from Crypto.Cipher import DES
import io
from Crypto.Util.Padding import pad, unpad
from docx.shared import Inches
import docx

def index(request):
    return render(request, "belge_yukleme.html")

def des_bilgilendirme(request):
    return render(request,"des_bilgi.html")

   
def belge_sifreleme(request):
    if request.method == 'POST' and request.FILES['belge']:
        # Word dosyasını yükle
        belge = request.FILES['belge']
        
        # Anahtar kelimeyi al
        anahtar_kelime = request.POST['anahtar_kelime']
        
        # Şifreleme için DES algoritmasını kullan
        cipher = DES.new(anahtar_kelime[:8].encode(), DES.MODE_ECB)

        # Word dosyasını oku
        doc = docx.Document(belge)
        
        # Yeni bir word dosyası oluştur
        yeni_belge = docx.Document()
        
        # Word dosyasındaki her bir paragrafı şifrele ve yeni belgeye ekle
        for para in doc.paragraphs:
            sifreli_metin = cipher.encrypt(pad(para.text.encode(), 8))
            yeni_para = yeni_belge.add_paragraph()
            yeni_para.add_run(sifreli_metin.hex())
        
        # Yeni belgeyi bir io stream'ine yaz ve kullanıcıya indirme seçeneği sun
        stream = io.BytesIO()
        yeni_belge.save(stream)
        stream.seek(0)
        
        response = HttpResponse(stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=sifreli_belge.docx'
        
        return response

    else:
        return render(request, 'sifreli_belge_yukleme.html')
    



def belge_desifreleme(request):
    if request.method == 'POST' and request.FILES['belge']:
        # Word dosyasını yükle
        belge = request.FILES['belge']
        
        # Anahtar kelimeyi al
        anahtar_kelime = request.POST['anahtar_kelime']
        
        # Şifreleme için DES algoritmasını kullan
        cipher = DES.new(anahtar_kelime[:8].encode(), DES.MODE_ECB)

        # Word dosyasını oku
        doc = docx.Document(belge)
        
        # Yeni bir word dosyası oluştur
        yeni_belge = docx.Document()
        
        # Dosyanın sonunda anahtar kelimesini ara ve şifrelenmiş veriyi al
        sifreli_veri = bytes.fromhex(doc.paragraphs[-1].text)
        
        # Anahtar kelimesini dosyanın sonundan sil
        doc.paragraphs[-1].text = ''
        
        # Şifre çözme işlemini gerçekleştir ve yeni belgeye ekle
        sifresiz_metin = unpad(cipher.decrypt(sifreli_veri), 8).decode()
        yeni_para = yeni_belge.add_paragraph()
        yeni_para.add_run(sifresiz_metin)
        
        # Diğer paragrafları şifresiz şekilde yeni belgeye ekle
        for para in doc.paragraphs:
            if para.text != '':
                yeni_para = yeni_belge.add_paragraph()
                yeni_para.add_run(para.text)
        
        # Yeni belgeyi bir io stream'ine yaz ve kullanıcıya indirme seçeneği sun
        stream = io.BytesIO()
        yeni_belge.save(stream)
        stream.seek(0)
        
        response = HttpResponse(stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=sifresiz_belge.docx'
        
        return response

    else:
        return render(request, 'belge_yukleme.html')